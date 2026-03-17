"""
Resume parser service for extracting text from PDF and DOCX files.
"""

import io
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document


class ResumeParser:
    """Utility class for extracting text from resume files."""

    @staticmethod
    async def extract_text(file_content: bytes, mime_type: str) -> str:
        """
        Extract text from a file based on its MIME type.

        Args:
            file_content: Raw bytes of the file.
            mime_type: MIME type of the file.

        Returns:
            Extracted text content.
        """
        if mime_type == "application/pdf":
            return ResumeParser._extract_from_pdf(file_content)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return ResumeParser._extract_from_docx(file_content)
        elif mime_type.startswith("text/"):
            return file_content.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to read PDF file: the file may be corrupted or malformed. Details: {e}")
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)

    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: the file may be corrupted or malformed. Details: {e}")
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)

    @staticmethod
    def parse_sections(text: str) -> dict:
        """
        Parse extracted text into structured sections.

        Returns a dict with keys like 'contact', 'summary', 'experience',
        'education', 'skills', etc.
        """
        sections = {
            "raw_text": text,
            "contact": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "certifications": "",
        }

        section_keywords = {
            "summary": ["summary", "objective", "about", "profile"],
            "experience": ["experience", "work history", "employment", "professional experience"],
            "education": ["education", "academic", "qualification"],
            "skills": ["skills", "technical skills", "competencies", "technologies"],
            "certifications": ["certifications", "certificates", "licenses"],
        }

        lines = text.split("\n")
        current_section = "contact"

        for line in lines:
            line_lower = line.strip().lower()
            matched = False
            for section, keywords in section_keywords.items():
                if any(kw in line_lower for kw in keywords) and len(line.strip()) < 50:
                    current_section = section
                    matched = True
                    break
            if not matched and line.strip():
                sections[current_section] += line + "\n"

        # Clean up sections
        for key in sections:
            if key != "raw_text":
                sections[key] = sections[key].strip()

        return sections

    @staticmethod
    def extract_skills_list(text: str) -> list[str]:
        """Extract a list of skills from resume text using keyword matching."""
        common_skills = [
            "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
            "react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi",
            "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
            "git", "ci/cd", "agile", "scrum", "rest api", "graphql",
            "machine learning", "deep learning", "nlp", "computer vision",
            "html", "css", "sass", "tailwind", "bootstrap",
            "linux", "bash", "powershell",
            "figma", "sketch", "adobe",
            "project management", "leadership", "communication",
        ]

        text_lower = text.lower()
        found_skills = []
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())

        return found_skills
